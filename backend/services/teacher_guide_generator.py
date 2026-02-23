"""Teacher Guide Generator - creates .docx teacher lesson guides"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.models.learning_experience import LearningExperience
from backend.models.lesson import Lesson
import json
import os

class TeacherGuideGenerator:
    """Generate comprehensive teacher guides as .docx files"""
    
    # KRPS Green brand color
    KRPS_GREEN = RGBColor(45, 139, 61)  # #2D8B3D
    
    @staticmethod
    def generate(lesson_id, output_dir='/tmp'):
        """
        Generate teacher guide for a lesson
        
        Args:
            lesson_id: ID of lesson
            output_dir: Where to save the file
        
        Returns:
            Dictionary with file_path and metadata
        """
        lesson = Lesson.query_by_id(lesson_id)
        if not lesson:
            return None
        
        le = LearningExperience.query_by_id(lesson.learning_experience_id)
        if not le:
            return None
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run('TEACHER LESSON GUIDE')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = TeacherGuideGenerator.KRPS_GREEN
        
        # Header info
        doc.add_paragraph(f"Subject: {le.subject}")
        doc.add_paragraph(f"Unit Number: Unit {le.unit_number}")
        doc.add_paragraph(f"Learning Experience: LE {le.experience_number}")
        doc.add_paragraph(f"Year Level: Year {le.year_level}")
        doc.add_paragraph(f"Duration: {lesson.duration_minutes} minutes")
        
        if le.nesa_outcome_code:
            doc.add_paragraph(f"NESA Outcome: {le.nesa_outcome_code}")
        
        doc.add_paragraph()  # Spacing
        
        # Core Concept
        heading = doc.add_paragraph()
        heading_run = heading.add_run('CORE CONCEPT')
        heading_run.font.bold = True
        heading_run.font.color.rgb = TeacherGuideGenerator.KRPS_GREEN
        doc.add_paragraph(le.core_concept)
        
        # Learning Intention
        heading = doc.add_paragraph()
        heading_run = heading.add_run('LEARNING INTENTION (WALT)')
        heading_run.font.bold = True
        heading_run.font.color.rgb = TeacherGuideGenerator.KRPS_GREEN
        doc.add_paragraph(le.learning_intention)
        
        # Success Criteria
        heading = doc.add_paragraph()
        heading_run = heading.add_run('SUCCESS CRITERIA (WILF)')
        heading_run.font.bold = True
        heading_run.font.color.rgb = TeacherGuideGenerator.KRPS_GREEN
        
        sc_list = le.get_success_criteria_list()
        for sc in sc_list:
            doc.add_paragraph(sc, style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # Lesson Structure
        heading = doc.add_paragraph()
        heading_run = heading.add_run('LESSON STRUCTURE: I DO, WE DO, YOU DO')
        heading_run.font.bold = True
        heading_run.font.color.rgb = TeacherGuideGenerator.KRPS_GREEN
        
        # I DO
        doc.add_paragraph('I DO (Teacher Demonstrates)', style='List Number')
        doc.add_paragraph('Duration: 15-20 minutes', style='List Bullet')
        doc.add_paragraph('Explicitly teach and model the concept', style='List Bullet')
        doc.add_paragraph('Use think-aloud to show your reasoning', style='List Bullet')
        doc.add_paragraph('Provide worked examples with annotations', style='List Bullet')
        
        # WE DO
        doc.add_paragraph('WE DO (Guided Practice)', style='List Number')
        doc.add_paragraph('Duration: 10-15 minutes', style='List Bullet')
        doc.add_paragraph('Work through 5 baseline questions together', style='List Bullet')
        doc.add_paragraph('Students attempt with teacher support', style='List Bullet')
        doc.add_paragraph('Discuss answers and address misconceptions', style='List Bullet')
        
        # YOU DO
        doc.add_paragraph('YOU DO (Independent Practice)', style='List Number')
        doc.add_paragraph('Duration: 15-20 minutes', style='List Bullet')
        doc.add_paragraph('Students complete differentiated worksheets', style='List Bullet')
        doc.add_paragraph('Four tiers: Mild (5Q), Medium (10Q), Spicy (15Q), Enrichment (1-2Q)', style='List Bullet')
        doc.add_paragraph('Monitor and provide individual support', style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # Teaching Notes
        heading = doc.add_paragraph()
        heading_run = heading.add_run('TEACHING NOTES')
        heading_run.font.bold = True
        heading_run.font.color.rgb = TeacherGuideGenerator.KRPS_GREEN
        
        doc.add_paragraph('Common misconceptions to watch for:')
        doc.add_paragraph('Use manipulatives and visual models to support understanding', style='List Bullet')
        doc.add_paragraph('Provide multiple explanations if students are confused', style='List Bullet')
        doc.add_paragraph('Celebrate effort and growth mindset', style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # Assessment
        heading = doc.add_paragraph()
        heading_run = heading.add_run('ASSESSMENT')
        heading_run.font.bold = True
        heading_run.font.color.rgb = TeacherGuideGenerator.KRPS_GREEN
        
        doc.add_paragraph('Formative Assessment:')
        doc.add_paragraph('Observe students during We Do phase', style='List Bullet')
        doc.add_paragraph('Use exit tickets to check understanding', style='List Bullet')
        doc.add_paragraph('Review student work samples from You Do', style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph('Use 4-level rubric:')
        doc.add_paragraph('1. Developing - Limited understanding', style='List Bullet')
        doc.add_paragraph('2. Approaching - Elementary understanding', style='List Bullet')
        doc.add_paragraph('3. Meeting - Adequate grade-level understanding', style='List Bullet')
        doc.add_paragraph('4. Exceeding - Sophisticated/extended understanding', style='List Bullet')
        
        # Footer with page numbers
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Page "
        footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save file
        filename = f"TeacherGuide_Unit{le.unit_number}_LE{le.experience_number}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        
        return {
            'file_path': filepath,
            'filename': filename,
            'lesson_id': lesson_id,
            'subject': le.subject,
            'core_concept': le.core_concept
        }
