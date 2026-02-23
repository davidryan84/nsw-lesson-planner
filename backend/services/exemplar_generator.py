"""Exemplar Generator - creates .docx exemplar work samples"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.models.learning_experience import LearningExperience
from backend.models.lesson import Lesson
from backend.models.worksheet import Worksheet
from backend.services.worksheet_service import WorksheetService
import os

class ExemplarGenerator:
    """Generate exemplar work samples showing quality responses"""
    
    KRPS_GREEN = RGBColor(45, 139, 61)
    
    @staticmethod
    def generate(lesson_id, output_dir='/tmp'):
        """
        Generate exemplar document showing model student work
        
        Args:
            lesson_id: ID of lesson
            output_dir: Where to save file
        
        Returns:
            Dictionary with file_path and metadata
        """
        lesson = Lesson.query_by_id(lesson_id)
        if not lesson:
            return None
        
        le = LearningExperience.query_by_id(lesson.learning_experience_id)
        if not le:
            return None
        
        # Get worksheets
        worksheets = WorksheetService.get_worksheets_by_lesson(lesson_id)
        if not worksheets:
            return None
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run('EXEMPLAR WORK - QUALITY RESPONSES')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = ExemplarGenerator.KRPS_GREEN
        
        doc.add_paragraph(f"Subject: {le.subject}")
        doc.add_paragraph(f"Topic: {le.core_concept}")
        doc.add_paragraph(f"Unit {le.unit_number} - Learning Experience {le.experience_number}")
        doc.add_paragraph()
        doc.add_paragraph("This document shows model responses at each tier level.")
        doc.add_paragraph()
        
        # Process each tier
        tier_order = ['mild', 'medium', 'spicy', 'enrichment']
        
        for tier in tier_order:
            ws = next((w for w in worksheets if w.tier == tier), None)
            if not ws:
                continue
            
            # Tier heading
            heading = doc.add_paragraph()
            heading_run = heading.add_run(f"{tier.upper()} TIER EXEMPLARS")
            heading_run.font.bold = True
            heading_run.font.color.rgb = ExemplarGenerator.KRPS_GREEN
            heading_run.font.size = Pt(13)
            
            doc.add_paragraph(f"This tier contains {ws.question_count} questions.")
            
            # Get questions
            questions = WorksheetService.get_questions(ws.id)
            
            # Show 1-3 exemplar questions per tier
            exemplar_count = min(3, len(questions))
            for i, q in enumerate(questions[:exemplar_count]):
                doc.add_paragraph()
                
                # Question
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Example {i+1} - Q{q.question_number}: ")
                q_run.bold = True
                q_para.add_run(q.question_text)
                
                # Model Response
                if q.model_answer:
                    resp_para = doc.add_paragraph()
                    resp_run = resp_para.add_run("MODEL RESPONSE: ")
                    resp_run.bold = True
                    resp_run.font.color.rgb = ExemplarGenerator.KRPS_GREEN
                    resp_para.add_run(q.model_answer)
                    
                    # Annotations
                    doc.add_paragraph("✓ Shows clear understanding", style='List Bullet')
                    doc.add_paragraph("✓ Includes appropriate working/reasoning", style='List Bullet')
                    
                    if tier == 'mild':
                        doc.add_paragraph("✓ Uses provided scaffolding effectively", style='List Bullet')
                    elif tier == 'spicy':
                        doc.add_paragraph("✓ Demonstrates extended thinking", style='List Bullet')
                    elif tier == 'enrichment':
                        doc.add_paragraph("✓ Applies learning to real-world context", style='List Bullet')
            
            # Assessment note
            doc.add_paragraph()
            assessment_para = doc.add_paragraph()
            assess_run = assessment_para.add_run("RUBRIC LEVEL: ")
            assess_run.bold = True
            
            if tier == 'mild':
                assessment_para.add_run("Approaching - Shows developing understanding with support")
            elif tier == 'medium':
                assessment_para.add_run("Meeting - Demonstrates grade-level competency")
            elif tier == 'spicy':
                assessment_para.add_run("Exceeding - Shows sophisticated reasoning")
            else:  # enrichment
                assessment_para.add_run("Exceeding - Extends learning beyond grade level")
            
            doc.add_paragraph()
            doc.add_page_break()
        
        # Save
        filename = f"Exemplar_Unit{le.unit_number}_LE{le.experience_number}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        
        return {
            'file_path': filepath,
            'filename': filename,
            'lesson_id': lesson_id,
            'subject': le.subject,
            'core_concept': le.core_concept
        }
