"""Answer Sheet Generator - creates .docx answer sheets for all question tiers"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.models.learning_experience import LearningExperience
from backend.models.lesson import Lesson
from backend.models.worksheet import Worksheet
from backend.services.worksheet_service import WorksheetService
import os

class AnswerSheetGenerator:
    """Generate comprehensive answer sheets as .docx files"""
    
    KRPS_GREEN = RGBColor(45, 139, 61)
    
    @staticmethod
    def generate(lesson_id, output_dir='/tmp'):
        """
        Generate answer sheet for all worksheet tiers
        
        Args:
            lesson_id: ID of lesson
            output_dir: Where to save file
        
        Returns:
            Dictionary with file_path and metadata
        """
        lesson = Lesson.query_by_id(lesson_id)
        if not lesson:
            return None
        
        le = LearningExperience.query_by_id(lesson.learning_experience_id)
        if not le:
            return None
        
        # Get worksheets
        worksheets = WorksheetService.get_worksheets_by_lesson(lesson_id)
        if not worksheets:
            return None
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run('ANSWER SHEET')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = AnswerSheetGenerator.KRPS_GREEN
        
        doc.add_paragraph(f"Subject: {le.subject}")
        doc.add_paragraph(f"Topic: {le.core_concept}")
        doc.add_paragraph(f"Unit {le.unit_number} - Learning Experience {le.experience_number}")
        doc.add_paragraph()
        
        # Process each worksheet tier
        for ws in worksheets:
            # Tier heading
            heading = doc.add_paragraph()
            heading_run = heading.add_run(f"{ws.tier.upper()} TIER ({ws.question_count} questions)")
            heading_run.font.bold = True
            heading_run.font.color.rgb = AnswerSheetGenerator.KRPS_GREEN
            heading_run.font.size = Pt(12)
            
            # Get questions
            questions = WorksheetService.get_questions(ws.id)
            
            for q in questions:
                # Question
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Q{q.question_number}: ")
                q_run.bold = True
                q_para.add_run(q.question_text)
                
                # Answer
                if q.model_answer:
                    ans_para = doc.add_paragraph()
                    ans_run = ans_para.add_run("Answer: ")
                    ans_run.italic = True
                    ans_para.add_run(q.model_answer)
                
                doc.add_paragraph()  # Spacing
            
            doc.add_paragraph()  # Tier spacing
        
        # Save
        filename = f"AnswerSheet_Unit{le.unit_number}_LE{le.experience_number}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        
        return {
            'file_path': filepath,
            'filename': filename,
            'lesson_id': lesson_id,
            'subject': le.subject,
            'core_concept': le.core_concept,
            'total_questions': sum(ws.question_count for ws in worksheets)
        }
